import gradio as gr
from groq import Groq
import os
import csv
import docx
import openpyxl

try:
    from pypdf import PdfReader
except ImportError:
    from PyPDF2 import PdfReader

# ── Groq client ───────────────────────────────────────────────────────────────
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# ── Extraction ────────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            parts.append(f"[Page {i+1}]\n{text.strip()}")
    return "\n\n".join(parts)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def extract_text_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)

def extract_text_from_csv(file_path):
    parts = []
    with open(file_path, newline="", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = " | ".join(cell.strip() for cell in row if cell.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)

def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()

def extract_document_text(file):
    if file is None:
        return None, "No file uploaded."
    path = file.name
    ext = os.path.splitext(path)[-1].lower()
    try:
        if ext == ".pdf":
            text = extract_text_from_pdf(path)
            label = "PDF Document"
        elif ext == ".docx":
            text = extract_text_from_docx(path)
            label = "Word Document (.docx)"
        elif ext in (".xlsx", ".xls"):
            text = extract_text_from_xlsx(path)
            label = "Excel Spreadsheet"
        elif ext == ".csv":
            text = extract_text_from_csv(path)
            label = "CSV File"
        elif ext in (".txt", ".md"):
            text = extract_text_from_txt(path)
            label = "Text File"
        else:
            return None, f"Unsupported file type: '{ext}'. Please upload PDF, DOCX, XLSX, CSV, or TXT."

        if not text.strip():
            return None, "Document appears to be empty or unreadable."

        word_count = len(text.split())
        filename = os.path.basename(path)
        status = (
            f"✅ '{filename}' loaded successfully.\n"
            f"📄 Type: {label}\n"
            f"📊 ~{word_count:,} words extracted\n"
            f"🤖 Ready to answer your mortgage questions!"
        )
        return text, status
    except Exception as e:
        return None, f"Error reading file: {str(e)}"

# ── AI Chat ───────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert mortgage analyst assistant. You have been given the full text of a mortgage-related document.

Your job is to:
- Answer any question about the document clearly and accurately
- Extract specific numbers, dates, names, fees, rates, and terms when asked
- Summarize sections if requested
- Handle all mortgage document types: fee worksheets, title reports, payslips, employment contracts, mortgage notes, deeds of trust, preliminary reports, affidavits, and more
- If the answer is not in the document, say so clearly — do NOT make up numbers or facts

Always be concise but thorough. Use bullet points for lists. Format dollar amounts and percentages clearly."""

def chat_with_doc(message, chat_history, doc_text):
    if not message.strip():
        return "", chat_history

    if client is None:
        chat_history.append({"role": "user", "content": message})
        chat_history.append({
            "role": "assistant",
            "content": "⚠️ No API key found. Add GROQ_API_KEY to your Hugging Face Space secrets and restart."
        })
        return "", chat_history

    if not doc_text:
        chat_history.append({"role": "user", "content": message})
        chat_history.append({
            "role": "assistant",
            "content": "Please upload and analyze a mortgage document first, then ask your question."
        })
        return "", chat_history

    messages = [{"role": "system", "content": SYSTEM_PROMPT}]
    messages.append({
        "role": "user",
        "content": f"Here is the mortgage document content:\n---\n{doc_text[:15000]}\n---"
    })
    messages.append({
        "role": "assistant",
        "content": "I have read the document. What would you like to know?"
    })
    for msg in chat_history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({"role": "user", "content": message})

    chat_history.append({"role": "user", "content": message})

    try:
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=messages,
            max_tokens=1000
        )
        answer = response.choices[0].message.content
    except Exception as e:
        answer = f"❌ Error: {str(e)}"

    chat_history.append({"role": "assistant", "content": answer})
    return "", chat_history

# ── CSS — matches photo layout, purple theme ──────────────────────────────────

css = """
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&display=swap');

body, .gradio-container {
    background: #0f0a1e !important;
    font-family: 'DM Sans', sans-serif !important;
    color: #ede9fe !important;
}
.gradio-container {
    max-width: 1200px !important;
    margin: 0 auto !important;
    padding: 24px !important;
}

/* Header — matches photo */
#app-header { margin-bottom: 24px; }
#app-header h1 {
    font-size: 1.5rem !important;
    font-weight: 700 !important;
    color: #ede9fe !important;
    margin: 0 0 4px 0 !important;
}
#app-header p {
    color: #a78bfa !important;
    font-size: 0.85rem !important;
    margin: 0 !important;
}

/* Section labels — matches photo */
.label-text {
    font-size: 13px !important;
    font-weight: 600 !important;
    color: #c4b5fd !important;
    margin-bottom: 4px;
    display: block;
}

/* Format badges row — matches photo */
.badge-row {
    font-size: 12px;
    color: #7c5cbf;
    margin-bottom: 10px;
    letter-spacing: 0.5px;
}

/* File upload — matches photo */
#file-upload {
    background: #1a1030 !important;
    border: 1.5px dashed #3d2a6e !important;
    border-radius: 10px !important;
    min-height: 200px !important;
}
#file-upload:hover { border-color: #7c3aed !important; }

/* Analyze button — full width, matches photo */
#analyze-btn {
    background: #3b1f6e !important;
    color: #ede9fe !important;
    border: none !important;
    border-radius: 8px !important;
    font-weight: 700 !important;
    font-size: 15px !important;
    height: 48px !important;
    width: 100% !important;
    margin-top: 12px !important;
}
#analyze-btn:hover { background: #7c3aed !important; }

/* Status label and textbox */
#status-label {
    font-size: 13px;
    font-weight: 600;
    color: #c4b5fd;
    margin-top: 16px;
    margin-bottom: 4px;
    display: block;
}
#status-box {
    margin-top: 4px;
}
#status-box textarea {
    background: #1a1030 !important;
    color: #ede9fe !important;
    border: 1px solid #3d2a6e !important;
    border-radius: 8px !important;
    font-size: 13px !important;
}

/* Chat label */
#chat-label {
    font-size: 13px;
    font-weight: 600;
    color: #c4b5fd;
    margin-bottom: 6px;
    display: block;
}

/* Chatbot window — matches photo */
#chat-window {
    background: #1a1030 !important;
    border: 1px solid #3d2a6e !important;
    border-radius: 10px !important;
    height: 460px !important;
}

/* Chat input — matches photo */
#chat-input textarea {
    background: #1a1030 !important;
    color: #ede9fe !important;
    border: 1px solid #3d2a6e !important;
    border-radius: 8px !important;
    font-size: 14px !important;
}
#chat-input textarea:focus { border-color: #7c3aed !important; }
#chat-input textarea::placeholder { color: #6d4aaa !important; }

/* Clear button — matches photo */
#clear-btn {
    background: #241545 !important;
    color: #a78bfa !important;
    border: 1px solid #3d2a6e !important;
    border-radius: 8px !important;
    font-weight: 600 !important;
    height: 46px !important;
}
#clear-btn:hover {
    background: #3d2a6e !important;
    color: #ede9fe !important;
}

footer { display: none !important; }
"""

# ── UI — matches photo layout exactly ────────────────────────────────────────

with gr.Blocks(css=css) as demo:

    doc_state = gr.State(None)

    # Header
    gr.HTML("""
        <div id="app-header">
            <h1>🏦 Mortgage Assistant</h1>
            <p>Upload any mortgage document and ask questions</p>
        </div>
    """)

    with gr.Row(equal_height=True):

        # ── Left column ───────────────────────────────────────────────────────
        with gr.Column(scale=1):

            gr.HTML("""
                <span class="label-text">Upload Document</span>
                <div class="badge-row">PDF &nbsp; DOCX &nbsp; XLSX &nbsp; CSV &nbsp; TXT</div>
            """)

            file_input = gr.File(
                label="",
                file_types=[".pdf", ".docx", ".xlsx", ".xls", ".csv", ".txt", ".md"],
                file_count="single",
                elem_id="file-upload"
            )

            analyze_btn = gr.Button("Analyze Document", elem_id="analyze-btn")

            gr.HTML("""<span id="status-label">Status</span>""")
            status_box = gr.Textbox(
                value="No document loaded. Upload a file to get started.",
                label="",
                elem_id="status-box",
                lines=4,
                interactive=False
            )

        # ── Right column ──────────────────────────────────────────────────────
        with gr.Column(scale=2):

            gr.HTML("""<span id="chat-label">Chat with your document</span>""")
            chatbot = gr.Chatbot(
                label="",
                elem_id="chat-window",
                height=460
            )

            with gr.Row():
                msg_input = gr.Textbox(
                    label="",
                    placeholder="e.g. What is the interest rate? Who is the borrower? What are the total fees?",
                    elem_id="chat-input",
                    scale=5,
                    container=False
                )
                clear_btn = gr.Button("Clear", elem_id="clear-btn", scale=1)

    # Events
    analyze_btn.click(fn=extract_document_text, inputs=file_input, outputs=[doc_state, status_box])
    msg_input.submit(fn=chat_with_doc, inputs=[msg_input, chatbot, doc_state], outputs=[msg_input, chatbot])
    clear_btn.click(fn=lambda: ([], ""), inputs=None, outputs=[chatbot, msg_input], queue=False)

if __name__ == "__main__":
    demo.launch(debug=True)
